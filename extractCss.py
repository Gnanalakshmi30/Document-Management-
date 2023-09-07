from docx import Document
from lxml import etree
doc = Document('C:/Users/New/Downloads/New folder/task.docx')
html = etree.tostring(doc._element.body).decode('utf-8')
styles = doc.styles
css = ''
for style in styles:
    css += f'{style.name} {{\n'
    for name, value in style.font._element.items():
        css += f'  {name}: {value};\n'
    css += '}\n'
content = ''
for paragraph in doc.paragraphs:
    content += f'<p>{paragraph.text}</p>\n'
for table in doc.tables:
    content += '<table>\n'
    for row in table.rows:
        content += '<tr>\n'
        for cell in row.cells:
            content += f'<td>{cell.text}</td>\n'
        content += '</tr>\n'
    content += '</table>\n'
final_html = f'<!DOCTYPE html>\n<html>\n<head>\n<style>\n{css}\n</style>\n</head>\n<body>\n{html}{content}</body>\n</html>'
with open('C:/Users/New/Downloads/New folder/example.html', 'w') as f:
    f.write(final_html)



