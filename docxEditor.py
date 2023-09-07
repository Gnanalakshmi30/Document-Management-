import docx
import sys
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import traceback



def replace_keyword(templatePath, keyvalPath):
    try:
        doc = docx.Document(templatePath)
        content = ''
        keyValues = open(keyvalPath, 'r')
        content = keyValues.read()
        keyValues.close()

        valueslst = content.split(',')
        

        if len(valueslst) > 0:
            for i in valueslst:
               
                keyval = i.split('|')
                key = '#{' + keyval[0] + '}#'
                val = keyval[1]
                
             
                # Replace the text in the body
                for para in doc.paragraphs:
                    if key in para.text:
                        if 'image' not in key.lower():
                            for run in para.runs:
                                if key in run.text:
                                    text = run.text.replace(
                                        key, val)
                                    padding = ' ' * (len(run.text) - len(text))
                                    font = run.font
                                    new_run = para.add_run()
                                    new_run.font.name = font.name
                                    new_run.font.size = font.size
                                    new_run.bold = font.bold
                                    new_run.italic = font.italic
                                    new_run.underline = font.underline
                                    new_run.subscript = font.subscript
                                    new_run.superscript = font.superscript
                                    new_run.font.color.rgb = font.color.rgb
                                    run.text = text
                        else:
                            image_paras = [i for i, p in enumerate(doc.paragraphs) if key in p.text]
                            p = doc.paragraphs[image_paras[0]]
                            p.alignment = 1
                            p.text = ""
                            r = p.add_run()
                            if val != '':
                                r.add_picture(val, width = Inches(5), height = Inches(2))
                            r.add_break()
                if 'result' in key.lower():
                    tables = doc.tables
                    for table in tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    if key in paragraph.text:
                                        inline = paragraph.runs
                                        for i in range(len(inline)):
                                            if key in inline[i].text:
                                                inline[i].text = inline[i].text.replace(key, val)

                if len(doc.sections) > 0:
                    # Replace the text in the header
                    if doc.sections[0].header is not None:
                        header = doc.sections[0].header
                        for header_paragraph in header.paragraphs:
                            if key in header_paragraph.text:
                                for run in header_paragraph.runs:
                                    if key in run.text:
                                        text = run.text.replace(
                                            key, val)
                                        padding = ' ' * (len(run.text) - len(text))
                                        font = run.font
                                        new_run = header_paragraph.add_run()
                                        new_run.font.name = font.name
                                        new_run.font.size = font.size
                                        new_run.bold = font.bold
                                        new_run.italic = font.italic
                                        new_run.underline = font.underline
                                        new_run.subscript = font.subscript
                                        new_run.superscript = font.superscript
                                        new_run.font.color.rgb = font.color.rgb
                                        run.text = text

                    # Replace the text in the footer
                    if doc.sections[0].footer is not None:
                        footer = doc.sections[0].footer
                        for paragraph in footer.paragraphs:
                            if key in paragraph.text:
                                for run in paragraph.runs:
                                    if key in run.text:
                                        text = run.text.replace(
                                            key, val)
                                        padding = ' ' * (len(run.text) - len(text))
                                        font = run.font
                                        new_run = paragraph.add_run()
                                        new_run.font.name = font.name
                                        new_run.font.size = font.size
                                        new_run.bold = font.bold
                                        new_run.italic = font.italic
                                        new_run.underline = font.underline
                                        new_run.subscript = font.subscript
                                        new_run.superscript = font.superscript
                                        new_run.font.color.rgb = font.color.rgb
                                        run.text = text

                doc.save(templatePath)
    except Exception as e:
        with open("log.txt", "w") as log:

           traceback.print_exc(file=log)
            
        print(e)


templatePath = sys.argv[1]
keyvalPath = sys.argv[2]
# templatePath = "C:\\Users\\vaanam EMP\\OneDrive - VAANAM TECHNOLOGIES PRIVATE LIMITED\\Documents\\PhotoApp\\10002\\GeneratedReport\\ELV Test Report Format-Metal.docx"
# keyvalPath = "C:\\Users\\vaanam EMP\\OneDrive - VAANAM TECHNOLOGIES PRIVATE LIMITED\\Documents\\PhotoApp\\10002\GeneratedReport\\ELV Test Report Format-Metal keyword.txt"

replace_keyword(templatePath, keyvalPath)
