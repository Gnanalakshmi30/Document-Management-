# import docx
# from bs4 import BeautifulSoup
# import sys

# def my_functionn(docxPath,htmlString):
#     try:
#         # Load the Word document
#         doc = docx.Document(docxPath)

#         # Get the first table in the document
#         table = doc.tables[0]

#         # Load the HTML table as a BeautifulSoup object
#         html = htmlString
#         soup = BeautifulSoup(html, 'html.parser')
#         html_table = soup.table

#         # Convert the HTML table to a list of lists
#         table_data = []
#         for row in html_table.find_all('tr'):
#             row_data = []
#             for cell in row.find_all('td'):
#                 row_data.append(cell.text)
#             table_data.append(row_data)

#         # Replace the contents of the first table with the HTML table data
#         for i, row in enumerate(table_data):
#             for j, cell in enumerate(row):
#                 table.cell(i, j).text = cell

#                  # Set the font family and size for the cell
#                 table.cell(i, j)._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), 'Arial')
#                 table.cell(i, j).paragraphs[0].runs[0].font.size = docx.shared.Pt(10)

#         # Save the modified document
#         doc.save(docxPath)
#     except Exception as e:
#         print(e)
#         return f"An error occurred: {str(e)}"  

# arg1 = sys.argv[1]
# arg2 = sys.argv[2]


# my_functionn(arg1,arg2)


import docx
from bs4 import BeautifulSoup
import sys

def my_functionn(docxPath,htmlString):
    try:
        # Load the Word document
        doc = docx.Document(docxPath)

        # Get the first table in the document
        table = doc.tables[0]

        # Load the HTML table as a BeautifulSoup object
        html = htmlString
        soup = BeautifulSoup(html, 'html.parser')
        html_table = soup.table

        # Convert the HTML table to a list of lists
        table_data = []
        for row in html_table.find_all('tr'):
            row_data = []
            for cell in row.find_all('td'):
                row_data.append(cell.text)
            table_data.append(row_data)

        # Replace the contents of the first table with the HTML table data
        for i, row in enumerate(table_data):
            for j, cell in enumerate(row):
                table.cell(i, j).text = cell

                # Set the font family and size for the paragraph in the cell
                for paragraph in table.cell(i, j).paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = docx.shared.Pt(10)

        # Save the modified document
        doc.save(docxPath)
    except Exception as e:
        print(e)
        return f"An error occurred: {str(e)}"  

arg1 = sys.argv[1]
arg2 = sys.argv[2]

my_functionn(arg1,arg2)
