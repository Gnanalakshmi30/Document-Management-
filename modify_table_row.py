import sys
from docx import Document
from docx.shared import Pt
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH



def modify_row_content(templatePath,actualtemplatePath, tableNo, row, col, cellData):
    try:
        if(row >= 0 and col >= 0):
            doc = Document(templatePath)
            table = doc.tables[tableNo]
            cell = table.cell(row, col)
            cell_text = cell.paragraphs[0].text
            cell.paragraphs[0].text = cellData
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(10)
            cell_font.name = 'Arial'
            doc.save(templatePath)

            doc2 = Document(actualtemplatePath)
            table = doc2.tables[tableNo]
            cell = table.cell(row, col)
            cell_text = cell.paragraphs[0].text
            cell.paragraphs[0].text = cellData
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(10)
            cell_font.name = 'Arial'
            doc2.save(actualtemplatePath)
    except Exception as e:
        print(e)
        return f"An error occurred: {str(e)}" 

   

templatePath = sys.argv[1]
actualtemplatePath = sys.argv[2]
tableNo = sys.argv[3]
row = sys.argv[4]
col = sys.argv[5]
cellData = sys.argv[6]

modify_row_content(templatePath,actualtemplatePath, int(tableNo), int(row), int(col), cellData)